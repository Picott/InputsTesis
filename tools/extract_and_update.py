import re
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document

# Paths
base = Path(r"c:\Users\juand\Desktop\TESIS FINAL SEGMENT\InputsTesis-main\InputsTesis-main")
docs = base / 'Documentation'
pdf1 = docs / 'ANEXO A.pdf'
pdf2 = docs / 'Avance_1_Juan David Basto Picott.pdf'
docx_path = docs / 'Informe_Final_Tesis.docx'

# Helper to extract text from PDF
def extract_text(pdf_path):
    text = []
    try:
        reader = PdfReader(str(pdf_path))
        for p in reader.pages:
            try:
                page_text = p.extract_text() or ''
            except Exception:
                page_text = ''
            text.append(page_text)
    except Exception as e:
        print(f"Error reading {pdf_path}: {e}")
    return "\n".join(text)

# Simple parsers
def find_university(text):
    m = re.search(r"Universid[aá]d\s*[:\-]?\s*([A-ZÁÉÍÓÚÑ][A-Za-zÁÉÍÓÚáéíóúñ ,.&-]{3,200})", text)
    if m:
        return m.group(1).strip()
    # fallback search for 'Universidad' line
    for line in text.splitlines():
        if 'Universidad' in line:
            return line.strip()
    return None

def find_director(text):
    m = re.search(r"(Director|Tutor|Director de tesis|Tutor de tesis)\s*[:\-]?\s*([A-Z][A-Za-zÁÉÍÓÚáéíóúñ ,.-]{3,100})", text, re.IGNORECASE)
    if m:
        return m.group(2).strip()
    return None

def find_date(text):
    # look for formats like 'Diciembre 2025' or '12 de Diciembre de 2025'
    m = re.search(r"(enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre)\s+\d{4}", text, re.IGNORECASE)
    if m:
        return m.group(0).strip().capitalize()
    m2 = re.search(r"\b(\d{1,2}\s+de\s+(enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre)\s+de\s+\d{4})\b", text, re.IGNORECASE)
    if m2:
        return m2.group(1).strip()
    return None

def find_dataset_info(text):
    # look for 'n = 123' or 'n=123' or 'muestras' or 'dataset'
    m = re.search(r"n\s*[=:]\s*(\d{2,7})", text, re.IGNORECASE)
    if m:
        return f"n = {m.group(1)}"
    m2 = re.search(r"(muestras|ejemplos|casos)\s*[:\-]?\s*(\d{2,7})", text, re.IGNORECASE)
    if m2:
        return f"{m2.group(1)} = {m2.group(2)}"
    m3 = re.search(r"dataset[s]?\s*[:\-]?\s*([\w\-_/ ]{3,80})", text, re.IGNORECASE)
    if m3:
        return m3.group(0)
    return None

def find_metrics(text):
    # Search for common metric words followed by numbers or percentages
    metrics = {}
    for key in ['Precisi[oó]n','Precisión','Accuracy','Recall','Sensibilidad','F1','IoU','AUC']:
        for m in re.finditer(rf"({key})\s*[:=]?\s*(\d+\.?\d*%?)", text, re.IGNORECASE):
            metrics[m.group(1).strip()] = m.group(2)
    # Also search for percentages standalone with metric nearby
    if not metrics:
        for m in re.finditer(r"(\d{1,3}\.\d+%|\d{1,3}%)", text):
            # capture context
            span = m.span()
            start = max(0, span[0]-40)
            ctx = text[start:span[1]+40]
            # find metric word in context
            if any(w in ctx.lower() for w in ['precis','recall','f1','iou','sensib','auc','accuracy']):
                metrics[ctx.strip()[:40]] = m.group(1)
    return metrics


# Main
text1 = extract_text(pdf1)
text2 = extract_text(pdf2)
combined = text1 + '\n' + text2

univ = find_university(combined) or 'UNIVERSIDAD - por completar'
director = find_director(combined) or 'Director - por completar'
date_found = find_date(combined) or ''
dataset = find_dataset_info(combined) or 'Datos: por completar'
metrics = find_metrics(combined)

print('University:', univ)
print('Director:', director)
print('Date:', date_found)
print('Dataset info:', dataset)
print('Metrics found:', metrics)

# Update docx
if docx_path.exists():
    doc = Document(str(docx_path))
    # Replace placeholders in paragraphs
    placeholders = {
        '[Nombre de la Universidad]': univ,
        '[Nombre del tutor]': director,
        '[Mes, Año]': date_found if date_found else '[Mes, Año]'
    }
    for p in doc.paragraphs:
        for ph, val in placeholders.items():
            if ph in p.text:
                p.text = p.text.replace(ph, val)
    # Find '6. RESULTADOS' or '6. Resultados' heading and insert metrics after
    inserted = False
    for i, p in enumerate(doc.paragraphs):
        if '6. RESULTADOS' in p.text.upper() or p.text.strip().upper().startswith('6. RESULTADOS') or p.text.strip().upper().startswith('6. RESULTADOS'):
            # insert after this paragraph
            insert_idx = i+1
            # Create content
            results_text = '\n'.join([f"{k}: {v}" for k, v in metrics.items()]) if metrics else 'Métricas: por completar'
            doc.paragraphs[i].insert_paragraph_before(results_text)
            inserted = True
            break
    if not inserted:
        # append at end
        doc.add_paragraph('\nRESULTADOS (Resumido extraído de anexos):')
        if metrics:
            for k, v in metrics.items():
                doc.add_paragraph(f"- {k}: {v}")
        else:
            doc.add_paragraph(dataset)
    # Save updated file
    updated_path = docs / 'Informe_Final_Tesis_actualizado.docx'
    doc.save(str(updated_path))
    print('Saved updated document at', updated_path)
else:
    print('Docx path not found:', docx_path)
