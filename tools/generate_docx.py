from docx import Document
from docx.shared import Pt

input_txt = r"c:\Users\juand\Desktop\TESIS FINAL SEGMENT\InputsTesis-main\InputsTesis-main\Documentation\Informe_Final_Tesis.txt"
output_docx = r"c:\Users\juand\Desktop\TESIS FINAL SEGMENT\InputsTesis-main\InputsTesis-main\Documentation\Informe_Final_Tesis.docx"

def add_paragraph_with_style(doc, text):
    # Simple heuristic: headings are lines in all-caps or starting with numbers
    stripped = text.strip()
    if not stripped:
        doc.add_paragraph('')
        return
    # Heading if line is short and mostly uppercase or starts with a number and a dot
    if len(stripped) < 80 and (stripped.isupper() or stripped[0].isdigit()):
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.bold = True
        run.font.size = Pt(12)
    else:
        p = doc.add_paragraph(stripped)
        p.style.font.size = Pt(11)


def main():
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    with open(input_txt, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        add_paragraph_with_style(doc, line)

    doc.save(output_docx)
    print(f"Saved {output_docx}")

if __name__ == '__main__':
    main()
